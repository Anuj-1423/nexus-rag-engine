"""
Document Parser Module
======================
Parses PDF, DOCX, and TXT files into a structured representation
with sections, headings, and metadata for downstream chunking.
"""

import io
import re
import os
import logging
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class Section:
    """A single section extracted from a document."""
    heading: str
    level: int              # 1 = top-level, 2 = sub-section, etc.
    content: str
    page_number: Optional[int] = None


@dataclass
class DocumentStructure:
    """Structured representation of a parsed document."""
    filename: str
    file_type: str          # "pdf", "docx", "txt"
    title: str              # Extracted or inferred document title
    sections: list[Section] = field(default_factory=list)
    raw_text: str = ""
    total_pages: Optional[int] = None


# ---------------------------------------------------------------------------
# PDF Parsing
# ---------------------------------------------------------------------------

def parse_pdf(file_bytes: bytes, filename: str) -> DocumentStructure:
    """
    Parse a PDF file into a DocumentStructure.
    
    Uses pypdf to extract text page-by-page. Attempts to detect headings
    by looking for short, uppercase or title-case lines that likely represent
    section titles.
    """
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    total_pages = len(reader.pages)

    all_text_parts: list[str] = []
    sections: list[Section] = []
    current_heading = "Introduction"
    current_content: list[str] = []
    current_page = 1

    for page_idx, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        all_text_parts.append(page_text)

        lines = page_text.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Heuristic: a line is likely a heading if it is short,
            # not ending with common punctuation, and is title-case or uppercase
            if _is_likely_heading(stripped):
                # Save the previous section
                if current_content:
                    sections.append(Section(
                        heading=current_heading,
                        level=_estimate_heading_level(current_heading),
                        content="\n".join(current_content).strip(),
                        page_number=current_page,
                    ))
                current_heading = stripped
                current_content = []
                current_page = page_idx
            else:
                current_content.append(stripped)

    # Flush last section
    if current_content:
        sections.append(Section(
            heading=current_heading,
            level=_estimate_heading_level(current_heading),
            content="\n".join(current_content).strip(),
            page_number=current_page,
        ))

    raw_text = "\n\n".join(all_text_parts).strip()

    # Safety Fallback: If no text was extracted but pages exist,
    # it might be a scanned PDF or pypdf failed.
    if not raw_text and total_pages > 0:
        logger.warning(f"No text extracted from {filename} using standard parser. PDF might be scanned.")
        raw_text = f"[Placeholder] This document '{filename}' appears to be an image-based or scanned PDF. Text extraction was unsuccessful."
        sections = [Section(heading="Unreadable Content", level=1, content=raw_text, page_number=1)]

    # Infer title from first heading or filename
    title = sections[0].heading if sections else _title_from_filename(filename)

    return DocumentStructure(
        filename=filename,
        file_type="pdf",
        title=title,
        sections=sections,
        raw_text=raw_text,
        total_pages=total_pages,
    )


# ---------------------------------------------------------------------------
# DOCX Parsing
# ---------------------------------------------------------------------------

def parse_docx(file_bytes: bytes, filename: str) -> DocumentStructure:
    """
    Parse a DOCX file into a DocumentStructure.
    
    Uses python-docx to extract paragraphs with their heading styles,
    which gives us reliable section hierarchy.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(file_bytes))

    sections: list[Section] = []
    current_heading = "Introduction"
    current_level = 1
    current_content: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()

        # Detect heading styles (Heading 1, Heading 2, etc.)
        if style_name.startswith("heading"):
            # Save previous section
            if current_content:
                sections.append(Section(
                    heading=current_heading,
                    level=current_level,
                    content="\n".join(current_content).strip(),
                ))
            # Extract heading level from style name
            level_match = re.search(r"(\d+)", style_name)
            current_level = int(level_match.group(1)) if level_match else 1
            current_heading = text
            current_content = []
        elif style_name == "title":
            if current_content:
                sections.append(Section(
                    heading=current_heading,
                    level=current_level,
                    content="\n".join(current_content).strip(),
                ))
            current_heading = text
            current_level = 1
            current_content = []
        else:
            current_content.append(text)

    # Also extract tables as text
    for table in doc.tables:
        table_text = _extract_table_text(table)
        if table_text:
            current_content.append(f"\n[Table]\n{table_text}")

    # Flush last section
    if current_content:
        sections.append(Section(
            heading=current_heading,
            level=current_level,
            content="\n".join(current_content).strip(),
        ))

    raw_text = "\n\n".join(
        s.content for s in sections
    )

    title = sections[0].heading if sections else _title_from_filename(filename)

    return DocumentStructure(
        filename=filename,
        file_type="docx",
        title=title,
        sections=sections,
        raw_text=raw_text,
    )


# ---------------------------------------------------------------------------
# TXT / Markdown Parsing
# ---------------------------------------------------------------------------

def parse_txt(file_bytes: bytes, filename: str) -> DocumentStructure:
    """
    Parse a plain text or markdown file into a DocumentStructure.
    
    Detects markdown-style headings (# H1, ## H2, etc.) and uses them
    to split into sections. Falls back to paragraph-based splitting.
    """
    text = file_bytes.decode("utf-8", errors="replace")
    lines = text.split("\n")

    sections: list[Section] = []
    current_heading = "Content"
    current_level = 1
    current_content: list[str] = []
    has_md_headings = False

    for line in lines:
        stripped = line.strip()

        # Detect markdown headings
        md_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if md_match:
            has_md_headings = True
            if current_content:
                sections.append(Section(
                    heading=current_heading,
                    level=current_level,
                    content="\n".join(current_content).strip(),
                ))
            current_level = len(md_match.group(1))
            current_heading = md_match.group(2).strip()
            current_content = []
        else:
            current_content.append(line)

    # Flush last section
    if current_content:
        sections.append(Section(
            heading=current_heading,
            level=current_level,
            content="\n".join(current_content).strip(),
        ))

    # If no markdown headings found, try splitting by blank-line groups
    if not has_md_headings and len(sections) == 1 and len(sections[0].content) > 2000:
        sections = _split_by_paragraphs(text, filename)

    title = sections[0].heading if sections else _title_from_filename(filename)

    return DocumentStructure(
        filename=filename,
        file_type="txt",
        title=title,
        sections=sections,
        raw_text=text,
    )


# ---------------------------------------------------------------------------
# Main Router
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}

def extract_document_structure(file_bytes: bytes, filename: str) -> DocumentStructure:
    """
    Route to the correct parser based on file extension.
    
    Args:
        file_bytes: Raw file content as bytes
        filename: Original filename with extension
        
    Returns:
        DocumentStructure with parsed sections and metadata
        
    Raises:
        ValueError: If file type is not supported
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return parse_pdf(file_bytes, filename)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_bytes, filename)
    elif ext in (".txt", ".md"):
        return parse_txt(file_bytes, filename)
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )


# ---------------------------------------------------------------------------
# Helper Functions
# ---------------------------------------------------------------------------

def _is_likely_heading(line: str) -> bool:
    """Heuristic to detect whether a line is likely a heading in a PDF."""
    # Too long → not a heading
    if len(line) > 100:
        return False
    # Too short → not meaningful
    if len(line) < 3:
        return False
    # Ends with sentence-ending punctuation → likely a sentence
    if line[-1] in ".,:;!?":
        return False
    # All caps or title case for lines under ~60 chars
    if len(line) <= 60 and (line.isupper() or line.istitle()):
        return True
    # Numbered section like "1.2 Overview" or "Section 3: Methods"
    if re.match(r"^(\d+\.?\d*\.?\d*)\s+\w", line):
        return True
    if re.match(r"^(section|chapter|part)\s+\d", line, re.IGNORECASE):
        return True
    return False


def _estimate_heading_level(heading: str) -> int:
    """Estimate heading depth from text patterns."""
    # "1.2.3 Title" → level 3
    num_match = re.match(r"^(\d+(?:\.\d+)*)", heading)
    if num_match:
        return num_match.group(1).count(".") + 1
    # All uppercase → level 1
    if heading.isupper():
        return 1
    return 2


def _title_from_filename(filename: str) -> str:
    """Derive a document title from the filename."""
    name = os.path.splitext(filename)[0]
    # Replace underscores/hyphens with spaces and title-case
    name = re.sub(r"[_\-]+", " ", name)
    return name.title()


def _extract_table_text(table) -> str:
    """Extract text from a python-docx table as a simple text representation."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows) if rows else ""


def _split_by_paragraphs(text: str, filename: str) -> list[Section]:
    """
    Fallback: split a long text into sections based on paragraph groups.
    Groups consecutive non-empty lines, separated by blank lines.
    """
    paragraphs = re.split(r"\n\s*\n", text)
    sections = []
    for i, para in enumerate(paragraphs):
        para = para.strip()
        if not para:
            continue
        # Use first line as heading (truncated)
        first_line = para.split("\n")[0][:80]
        sections.append(Section(
            heading=first_line if len(first_line) < 60 else f"Section {i + 1}",
            level=2,
            content=para,
        ))
    return sections if sections else [Section(
        heading=_title_from_filename(filename),
        level=1,
        content=text,
    )]
